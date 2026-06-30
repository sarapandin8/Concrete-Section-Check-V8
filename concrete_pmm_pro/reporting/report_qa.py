"""Automated QA checks for draft Word report content."""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

import pandas as pd
from docx import Document

from concrete_pmm_pro.reporting.report_manifest import ReportManifest


@dataclass(frozen=True)
class ReportQAItem:
    category: str
    check_name: str
    status: str
    message: str
    details: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ReportQASummary:
    overall_status: str
    pass_count: int
    warning_count: int
    fail_count: int
    items: list[ReportQAItem]
    warnings: list[str] = field(default_factory=list)
    info: list[str] = field(default_factory=list)


def report_qa_summary_to_dataframe(summary: ReportQASummary) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Category": item.category,
                "Check": item.check_name,
                "Status": item.status,
                "Message": item.message,
            }
            for item in summary.items
        ],
        columns=["Category", "Check", "Status", "Message"],
    )


def _document_from_bytes(docx_bytes: bytes):
    try:
        return Document(BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError("Invalid DOCX bytes; Word report QA could not read the document.") from exc


def extract_docx_text(docx_bytes: bytes) -> str:
    document = _document_from_bytes(docx_bytes)
    text: list[str] = []
    text.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text.append(cell.text)
    return "\n".join(text)


def extract_docx_headings(docx_bytes: bytes) -> list[str]:
    document = _document_from_bytes(docx_bytes)
    headings: list[str] = []
    known_headings = {
        "executive summary",
        "analysis mode and scope",
        "result traceability snapshot",
        "report readiness",
        "engineering warnings",
        "engineering limitations",
        "unit conventions",
        "standard terminology",
        "available report tables",
        "available report figures",
        "draft figures",
        "report generation notes",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith(("heading", "title")) or text.lower() in known_headings:
            headings.append(text)
    return headings


def _contains(text: str, phrase: str) -> bool:
    return phrase.casefold() in text.casefold()


def _qa_item(category: str, check_name: str, status: str, message: str, **details: Any) -> ReportQAItem:
    return ReportQAItem(category=category, check_name=check_name, status=status, message=message, details=details)


def validate_required_report_headings(docx_bytes: bytes) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    required = [
        "Executive Summary",
        "Analysis Mode and Scope",
        "Result Traceability Snapshot",
        "Report Readiness",
        "Engineering Warnings",
        "Engineering Limitations",
        "Unit Conventions",
        "Standard Terminology",
        "Available Report Tables",
        "Available Report Figures",
        "Draft Figures",
        "Report Generation Notes",
    ]
    items: list[ReportQAItem] = []
    for heading in required:
        if _contains(text, heading):
            items.append(_qa_item("Structure", heading, "PASS", f"Required section is present: {heading}"))
        else:
            items.append(_qa_item("Structure", heading, "FAIL", f"Required section is missing: {heading}"))
    return items


def validate_draft_disclaimer(docx_bytes: bytes) -> ReportQAItem:
    text = extract_docx_text(docx_bytes).casefold()
    has_draft = "draft engineering report" in text
    has_not_final = "not a final design certification" in text or "not final certification" in text
    has_current_results = (
        "generated from current analysis results" in text
        or "generated from current stored analysis results" in text
        or "stored session results" in text
    )
    if has_draft and has_not_final and has_current_results:
        return _qa_item("Disclaimer", "Draft/non-certification disclaimer", "PASS", "Draft and non-certification wording is present.")
    if has_draft or has_not_final:
        return _qa_item(
            "Disclaimer",
            "Draft/non-certification disclaimer",
            "WARNING",
            "Draft wording is present but the disclaimer is incomplete.",
            has_draft=has_draft,
            has_not_final=has_not_final,
            has_current_results=has_current_results,
        )
    return _qa_item(
        "Disclaimer",
        "Draft/non-certification disclaimer",
        "FAIL",
        "Report is missing clear draft/non-certification wording.",
    )


def _limitation_present(text: str, title: str, key: str) -> bool:
    key_phrase = key.replace("_", " ")
    return _contains(text, title) or _contains(text, key_phrase)


def validate_engineering_limitations_present(docx_bytes: bytes, manifest: ReportManifest) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    items: list[ReportQAItem] = []
    high_critical = [item for item in manifest.engineering_limitations if item.risk_level in {"HIGH", "CRITICAL"}]
    if _contains(text, "High/Critical limitation count") or _contains(text, "High/Critical limitations requiring review"):
        items.append(_qa_item("Limitations", "High/Critical limitation count", "PASS", "High/Critical limitation count or list is present."))
    else:
        items.append(_qa_item("Limitations", "High/Critical limitation count", "FAIL", "High/Critical limitation count/list is missing."))
    for limitation in manifest.engineering_limitations:
        present = _limitation_present(text, limitation.title, limitation.key)
        if present:
            items.append(
                _qa_item("Limitations", limitation.key, "PASS", f"Limitation is disclosed: {limitation.title}")
            )
        elif limitation.risk_level in {"HIGH", "CRITICAL"}:
            items.append(
                _qa_item("Limitations", limitation.key, "FAIL", f"HIGH/CRITICAL limitation is missing: {limitation.title}")
            )
        elif limitation.risk_level == "MEDIUM":
            items.append(_qa_item("Limitations", limitation.key, "WARNING", f"MEDIUM limitation is missing: {limitation.title}"))
    if not high_critical:
        items.append(_qa_item("Limitations", "High/Critical limitation disclosure", "PASS", "No HIGH/CRITICAL limitations are present in manifest."))
    return items


def validate_engineering_warnings_present(docx_bytes: bytes, manifest: ReportManifest) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    items = []
    if not _contains(text, "Engineering Warnings"):
        return [_qa_item("Warnings", "Engineering Warnings section", "FAIL", "Engineering Warnings section is missing.")]
    items.append(_qa_item("Warnings", "Engineering Warnings section", "PASS", "Engineering Warnings section is present."))
    if not manifest.engineering_warnings:
        status = "PASS" if _contains(text, "No engineering warnings recorded.") else "WARNING"
        message = "Report explicitly states no engineering warnings were recorded." if status == "PASS" else "Report has no warnings in manifest but does not state that clearly."
        items.append(_qa_item("Warnings", "Empty warning disclosure", status, message))
        return items
    found = [warning for warning in manifest.engineering_warnings if _contains(text, str(warning)[:80])]
    if found:
        items.append(_qa_item("Warnings", "Warning details", "PASS", f"{len(found)} engineering warning(s) appear in report."))
    else:
        items.append(_qa_item("Warnings", "Warning details", "FAIL", "Manifest warnings are not disclosed in report text."))
    return items


def validate_unit_conventions_present(docx_bytes: bytes) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    checks = {
        "Force / kN": ["Force", "kN"],
        "Moment / kN-m": ["Moment", "kN-m"],
        "Stress / MPa": ["Stress", "MPa"],
        "Length / mm": ["Length", "mm"],
        "Area / mm2": ["Area", "mm2"],
    }
    items = []
    if not _contains(text, "Unit Conventions"):
        return [_qa_item("Units", "Unit Conventions section", "WARNING", "Unit Conventions section is missing.")]
    for name, tokens in checks.items():
        status = "PASS" if all(_contains(text, token) for token in tokens) else "WARNING"
        message = f"{name} appears in unit conventions." if status == "PASS" else f"{name} was not found clearly."
        items.append(_qa_item("Units", name, status, message))
    return items


def validate_terminology_present(docx_bytes: bytes) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    terms = ["Pu", "Mux", "Muy", "Pe_eff", "No-Tension", "Decompression"]
    if not _contains(text, "Standard Terminology"):
        return [_qa_item("Terminology", "Standard Terminology section", "WARNING", "Standard Terminology section is missing.")]
    return [
        _qa_item(
            "Terminology",
            term,
            "PASS" if _contains(text, term) else "WARNING",
            f"{term} terminology is present." if _contains(text, term) else f"{term} terminology was not found clearly.",
        )
        for term in terms
    ]


def validate_traceability_and_readiness_present(docx_bytes: bytes, manifest: ReportManifest) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    checks = [
        ("Readiness status", manifest.readiness_summary.overall_status, "FAIL"),
        ("Generated status", manifest.generated_status, "WARNING"),
        ("Analysis mode / member type", manifest.traceability_snapshot.member_type or "Member type", "WARNING"),
        ("ULS result available", "ULS result available", "WARNING"),
        ("SLS result available", "SLS result available", "WARNING"),
        ("Warning count", "Warning count", "WARNING"),
        ("Limitation count", "High/Critical limitation count", "WARNING"),
    ]
    items: list[ReportQAItem] = []
    for check_name, token, missing_status in checks:
        if token and _contains(text, str(token)):
            items.append(_qa_item("Traceability", check_name, "PASS", f"{check_name} is present."))
        else:
            items.append(_qa_item("Traceability", check_name, missing_status, f"{check_name} is missing or unclear."))
    return items


def validate_report_tables_and_figures_present(docx_bytes: bytes, manifest: ReportManifest) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes)
    items = [
        _qa_item(
            "Tables/Figures",
            "Available Report Tables",
            "PASS" if _contains(text, "Available Report Tables") else "FAIL",
            "Available Report Tables section is present." if _contains(text, "Available Report Tables") else "Available Report Tables section is missing.",
        ),
        _qa_item(
            "Tables/Figures",
            "Available Report Figures",
            "PASS" if _contains(text, "Available Report Figures") else "FAIL",
            "Available Report Figures section is present." if _contains(text, "Available Report Figures") else "Available Report Figures section is missing.",
        ),
    ]
    export_ready = [item for item in manifest.figure_export_items if item.export_ready]
    if export_ready:
        has_placeholder = _contains(text, "Figure not embedded") or _contains(text, "Figure ")
        items.append(
            _qa_item(
                "Tables/Figures",
                "Figure embedding or placeholder",
                "PASS" if has_placeholder else "WARNING",
                "Figure embedding or placeholder text is present." if has_placeholder else "Export-ready figures exist but no embedding/placeholder text was found.",
            )
        )
    bar_caption_count = text.casefold().count("sls stress diagram")
    duplicate_visualization = _contains(text, "sls_stress_visualization") and bar_caption_count > 1
    items.append(
        _qa_item(
            "Tables/Figures",
            "Duplicate SLS bar chart captions",
            "WARNING" if duplicate_visualization else "PASS",
            "Duplicate SLS bar chart semantic captions may be present." if duplicate_visualization else "No duplicate SLS bar chart captions detected.",
        )
    )
    return items


def validate_no_misleading_certification_language(docx_bytes: bytes) -> list[ReportQAItem]:
    text = extract_docx_text(docx_bytes).casefold()
    problematic = [
        "certified design",
        "final approved design",
        "code-certified software",
        "guaranteed safe",
        "fully validated",
        "production-grade design certification",
    ]
    items: list[ReportQAItem] = []
    for phrase in problematic:
        index = text.find(phrase)
        if index < 0:
            items.append(_qa_item("Language", phrase, "PASS", f"Misleading phrase not found: {phrase}"))
            continue
        window = text[max(0, index - 80) : index + len(phrase) + 80]
        is_negated = any(
            token in window
            for token in ["not ", "no ", "future work", "prototype", "draft", "excluded", "engineering-review", "guarded", "checks"]
        )
        items.append(
            _qa_item(
                "Language",
                phrase,
                "PASS" if is_negated else "FAIL",
                f"Phrase is safely contextualized: {phrase}" if is_negated else f"Misleading certification language found: {phrase}",
            )
        )
    return items


def run_word_report_qa(docx_bytes: bytes, manifest: ReportManifest) -> ReportQASummary:
    items: list[ReportQAItem] = []
    items.extend(validate_required_report_headings(docx_bytes))
    items.append(validate_draft_disclaimer(docx_bytes))
    items.extend(validate_engineering_limitations_present(docx_bytes, manifest))
    items.extend(validate_engineering_warnings_present(docx_bytes, manifest))
    items.extend(validate_unit_conventions_present(docx_bytes))
    items.extend(validate_terminology_present(docx_bytes))
    items.extend(validate_traceability_and_readiness_present(docx_bytes, manifest))
    items.extend(validate_report_tables_and_figures_present(docx_bytes, manifest))
    items.extend(validate_no_misleading_certification_language(docx_bytes))

    pass_count = sum(1 for item in items if item.status == "PASS")
    warning_count = sum(1 for item in items if item.status == "WARNING")
    fail_count = sum(1 for item in items if item.status == "FAIL")
    overall_status = "FAIL" if fail_count else "WARNING" if warning_count else "PASS"
    return ReportQASummary(
        overall_status=overall_status,
        pass_count=pass_count,
        warning_count=warning_count,
        fail_count=fail_count,
        items=items,
        warnings=[item.message for item in items if item.status in {"WARNING", "FAIL"}],
        info=["Word report QA checks document content only; they do not rerun engineering solvers."],
    )
