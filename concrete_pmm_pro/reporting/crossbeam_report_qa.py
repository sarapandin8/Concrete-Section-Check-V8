"""Crossbeam Report / QA draft-export helpers.

This module is intentionally downstream-only: it accepts an already-normalized
report context and never prepares or runs Analysis solvers.  The Streamlit
workspace owns result selection/traceability; this module only formats the
stored evidence into a review draft.
"""

from __future__ import annotations

from io import BytesIO
from typing import Mapping, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _display(value: object) -> str:
    if value is None:
        return "-"
    text = str(value).strip()
    return text if text else "-"


def _set_cell_text(cell, value: object, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_display(value))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)


def _add_table(document: Document, rows: Sequence[Mapping[str, object]], columns: Sequence[str]) -> None:
    if not rows:
        document.add_paragraph("No stored evidence is available.")
        return
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader")
    header_repeat.set(qn("w:val"), "true")
    header_tr_pr.append(header_repeat)
    for index, column in enumerate(columns):
        _set_cell_text(table.rows[0].cells[index], column, bold=True)
    for row in rows:
        table_row = table.add_row()
        row_pr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        cells = table_row.cells
        for index, column in enumerate(columns):
            _set_cell_text(cells[index], row.get(column))


def _apply_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    for style_name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 11)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def _add_footer(document: Document) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.text = "Concrete Section Pro — Draft Design Report — engineering review only"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_crossbeam_draft_design_report(context: Mapping[str, object]) -> bytes:
    """Build a review-only Crossbeam design report DOCX from stored-result context.

    ``context`` must already contain stored-result summaries.  This function is
    deliberately solver-free and does not inspect or mutate Analysis inputs.
    """

    document = Document()
    _apply_document_defaults(document)
    _add_footer(document)

    title = _display(context.get("Report title"))
    document.add_heading(title if title != "-" else "Concrete Section Pro — Crossbeam Design Report", level=0)
    draft = document.add_paragraph()
    draft_run = draft.add_run("DRAFT — NOT FOR ISSUE")
    draft_run.bold = True
    draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Generated from stored Analysis results only. This draft does not rerun Flexure, Shear, Torsion, "
        "Combined V+T, SLS, or verification solvers. Final issue remains subject to Report / QA readiness."
    )

    document.add_heading("Report Metadata", level=1)
    metadata_rows = [
        {"Item": key, "Value": context.get(key)}
        for key in ("Project", "Prepared by", "Checked by", "Revision", "Construction type", "Design code", "Units")
    ]
    _add_table(document, metadata_rows, ("Item", "Value"))

    document.add_heading("Report Readiness", level=1)
    readiness_rows = [
        {"Item": "Overall status", "Value": context.get("Overall status")},
        {"Item": "Critical check", "Value": context.get("Critical check")},
        {"Item": "Report readiness", "Value": context.get("Report readiness")},
        {"Item": "ULS results", "Value": context.get("ULS completeness")},
        {"Item": "SLS status", "Value": context.get("SLS status")},
        {"Item": "Runtime mode", "Value": "READ-ONLY / stored results only"},
    ]
    _add_table(document, readiness_rows, ("Item", "Value"))
    readiness_note = _display(context.get("Readiness note"))
    if readiness_note != "-":
        document.add_paragraph(readiness_note)

    document.add_heading("Design Basis and Analysis Scope", level=1)
    design_basis_rows = list(context.get("Design basis rows") or [])
    _add_table(document, design_basis_rows, ("Item", "Value"))

    document.add_heading("Stored ULS Result Evidence", level=1)
    uls_rows = list(context.get("ULS rows") or [])
    document.add_paragraph("Governing decision summary")
    _add_table(
        document,
        uls_rows,
        (
            "Check",
            "Status",
            "Governing Check",
            "Station / Point",
            "D/C / Util.",
        ),
    )
    document.add_page_break()
    document.add_paragraph("Demand / capacity evidence and required action")
    _add_table(
        document,
        uls_rows,
        (
            "Check",
            "Demand",
            "Capacity / Limit",
            "Required Action",
        ),
    )

    document.add_heading("Required Actions", level=1)
    action_rows = list(context.get("Action rows") or [])
    _add_table(document, action_rows, ("Priority", "Module", "Issue", "Required Action"))

    document.add_heading("QA Scope Guards and Limitations", level=1)
    limitation_rows = list(context.get("Limitation rows") or [])
    _add_table(document, limitation_rows, ("Status", "Item", "Engineering meaning"))

    document.add_heading("Stored Result Traceability", level=1)
    trace_rows = list(context.get("Traceability rows") or [])
    _add_table(
        document,
        trace_rows,
        ("Check", "Status", "Construction type", "Case", "Station / Point", "Result fingerprint", "Source"),
    )

    document.add_paragraph(
        "Final Design Report export remains disabled until the application reports READY and all mandatory report gates are closed."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
