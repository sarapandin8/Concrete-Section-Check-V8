"""Draft Word report export from an existing report manifest."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Inches, Pt

from concrete_pmm_pro.reporting.figure_export import plotly_figure_to_png_bytes
from concrete_pmm_pro.reporting.limitations import engineering_limitations_to_dataframe
from concrete_pmm_pro.reporting.report_figures import build_exportable_figure, report_figure_export_items_to_dataframe
from concrete_pmm_pro.reporting.report_manifest import ReportManifest, report_manifest_to_summary_dataframe
from concrete_pmm_pro.reporting.report_sections import report_sections_to_dataframe
from concrete_pmm_pro.reporting.report_tables import report_tables_to_dataframe
from concrete_pmm_pro.reporting.readiness import report_readiness_to_dataframe
from concrete_pmm_pro.reporting.terminology import terminology_to_dataframe
from concrete_pmm_pro.reporting.traceability import result_traceability_snapshot_to_dataframe
from concrete_pmm_pro.reporting.units import unit_conventions_to_dataframe
from concrete_pmm_pro.reporting.railway_u_girder_report import build_railway_u_girder_sls_report_package
from concrete_pmm_pro.reporting.generic_precast_lifting_report import build_generic_precast_lifting_report_package
from concrete_pmm_pro.analysis.railway_u_girder_uls import build_railway_u_girder_uls_framework_package
from concrete_pmm_pro.reporting.railway_u_girder_release import build_railway_u_girder_release_package
from concrete_pmm_pro.reporting.railway_u_girder_final import build_railway_u_girder_final_design_check_package
from concrete_pmm_pro.reporting.column_pier_vt_report import build_column_pier_vt_report_package


@dataclass(frozen=True)
class ReportExportOptions:
    include_appendices: bool = True
    include_figures: bool = True
    max_table_rows: int = 30
    include_full_terminology: bool = True
    include_full_registries: bool = True


def add_report_heading(document: DocumentObject, text: str, level: int = 1) -> None:
    document.add_heading(text, level=max(1, min(level, 3)))


def apply_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)


def apply_heading_styles(document: DocumentObject) -> None:
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def set_table_style(table, compact: bool = True) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8 if compact else 9)


def add_page_break(document: DocumentObject) -> None:
    document.add_page_break()


def add_report_footer_note(document: DocumentObject) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.text = (
            "Concrete Section Pro draft report - engineering review; "
            "finalized production-preview only where explicitly validated"
        )
        footer.style = document.styles["Footer"]


def add_draft_disclaimer(document: DocumentObject) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "This document is a draft engineering report generated from current analysis results. "
        "It is not a final design certification."
    )
    run.bold = True


def add_key_value_table(document: DocumentObject, items, title: str | None = None) -> None:
    rows = [{"Item": key, "Value": value} for key, value in items]
    dataframe_to_word_table(document, pd.DataFrame(rows), title=title)


def _as_display(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def dataframe_to_word_table(
    document: DocumentObject,
    dataframe: pd.DataFrame,
    title: str | None = None,
    max_rows: int | None = None,
) -> None:
    if title:
        document.add_paragraph(title, style=None).runs[0].bold = True
    if dataframe is None or dataframe.empty:
        document.add_paragraph("No data available.")
        return
    source_rows = len(dataframe)
    display_df = dataframe.copy()
    if max_rows is not None and source_rows > max_rows:
        display_df = display_df.head(max_rows)
    table = document.add_table(rows=1, cols=len(display_df.columns))
    header_cells = table.rows[0].cells
    for index, column in enumerate(display_df.columns):
        header_cells[index].text = str(column)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(display_df.columns):
            cells[index].text = _as_display(row[column])
    set_table_style(table)
    if max_rows is not None and source_rows > max_rows:
        document.add_paragraph("Table truncated for draft report. Full data available from CSV export.")


def add_warning_box_or_section(document: DocumentObject, warnings: list[str], title: str) -> None:
    add_report_heading(document, title, level=2)
    if not warnings:
        document.add_paragraph("No engineering warnings recorded.")
        return
    for warning in warnings:
        document.add_paragraph(str(warning), style="List Bullet")


def add_limitations_section(document: DocumentObject, limitations) -> None:
    add_report_heading(document, "Engineering Limitations", level=1)
    ordered = sorted(
        list(limitations),
        key=lambda item: {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(item.risk_level, 4),
    )
    high_critical = [item for item in ordered if item.risk_level in {"HIGH", "CRITICAL"}]
    if high_critical:
        document.add_paragraph("High/Critical limitations requiring review:")
        for item in high_critical:
            document.add_paragraph(f"{item.title} ({item.risk_level})", style="List Bullet")
    dataframe = engineering_limitations_to_dataframe(ordered)
    keep_columns = ["Title", "Risk Level", "Status", "Category", "User Note", "Recommended Action"]
    dataframe_to_word_table(document, dataframe[keep_columns], max_rows=50)


def add_unit_conventions_section(document: DocumentObject, unit_conventions) -> None:
    add_report_heading(document, "Unit Conventions", level=1)
    dataframe = pd.DataFrame(
        [
            {
                "Quantity": item.quantity,
                "Internal Unit": item.internal_unit,
                "Display Unit": item.display_unit,
                "Report Unit": item.report_unit,
                "Conversion Note": item.conversion_note,
            }
            for item in unit_conventions
        ]
    )
    dataframe_to_word_table(document, dataframe if not dataframe.empty else unit_conventions_to_dataframe(), max_rows=50)


def add_terminology_section(document: DocumentObject, terminology) -> None:
    add_report_heading(document, "Standard Terminology", level=1)
    dataframe = pd.DataFrame(
        [
            {
                "Key": item.key,
                "Label": item.label,
                "Description": item.description,
                "Unit": item.unit or "",
                "Category": item.category,
            }
            for item in terminology
        ]
    )
    dataframe_to_word_table(document, dataframe if not dataframe.empty else terminology_to_dataframe(), max_rows=60)


def _add_cover(document: DocumentObject, manifest: ReportManifest) -> None:
    metadata = manifest.metadata
    document.add_heading(metadata.report_title, level=0)
    document.add_paragraph("Draft engineering report generated from current stored analysis results.")
    rows = [
        ("Project name", metadata.project_name),
        ("Organization", metadata.organization),
        ("Project number", metadata.project_number),
        ("Prepared by", metadata.prepared_by),
        ("Checked by", metadata.checked_by),
        ("Revision", metadata.revision),
        ("Report date", metadata.report_date),
        ("Draft status", "Draft / not final certification"),
        ("Note", metadata.note),
    ]
    add_key_value_table(document, rows)
    add_draft_disclaimer(document)
    add_page_break(document)


def _add_executive_summary(document: DocumentObject, manifest: ReportManifest) -> None:
    add_report_heading(document, "Executive Summary", level=1)
    snapshot = manifest.traceability_snapshot
    document.add_paragraph(
        "This is a draft engineering report generated from current stored analysis results. "
        "ACI RC Flexural PMM may be treated as finalized production-preview only within the validated RC scope; "
        "unsupported routes remain engineering-review items."
    )
    summary = {
        "Readiness status": manifest.readiness_summary.overall_status,
        "Analysis mode / member type": snapshot.member_type or "",
        "ULS result available": "Yes" if getattr(snapshot, "uls_result_available", False) else "No",
        "SLS result available": "Yes" if snapshot.sls_result_available else "No",
        "Governing ULS combo": snapshot.governing_uls_combo or "Not available",
        "Governing SLS combo": snapshot.governing_sls_combo or "Not available",
        "Governing SLS point": snapshot.governing_sls_point or "Not available",
        "Warning count": len(manifest.engineering_warnings),
        "High/Critical limitation count": len(
            [item for item in manifest.engineering_limitations if item.risk_level in {"HIGH", "CRITICAL"}]
        ),
        "Available figure count": len([figure for figure in manifest.figures if figure.available]),
    }
    dataframe_to_word_table(document, pd.DataFrame([{"Item": key, "Value": value} for key, value in summary.items()]))
    if not getattr(snapshot, "uls_result_available", False):
        document.add_paragraph("No stored ULS result is currently available in the stored session.")
    if not snapshot.sls_result_available:
        document.add_paragraph("No SLS result is currently available in the stored session.")


def _add_analysis_scope(document: DocumentObject, manifest: ReportManifest) -> None:
    add_report_heading(document, "Analysis Mode and Scope", level=1)
    snapshot = manifest.traceability_snapshot
    scope_rows = [
        {"Item": "Member type", "Value": snapshot.member_type or ""},
        {"Item": "Analysis workflow", "Value": snapshot.analysis_workflow or ""},
        {"Item": "ULS workflow status", "Value": "Available" if getattr(snapshot, "uls_result_available", False) else "Not run / not available"},
        {"Item": "SLS workflow status", "Value": "Available" if snapshot.sls_result_available else "Not run / not available"},
    ]
    dataframe_to_word_table(document, pd.DataFrame(scope_rows))
    if snapshot.member_type == "beam_girder":
        document.add_paragraph("Beam/Girder ULS/SLS preview checks are guarded engineering-review outputs. Final code-certified girder design, development length, anchorage, end-zone, interface shear, fatigue, and seismic/detailing certification remain outside this draft report scope.")


def _add_figures(document: DocumentObject, manifest: ReportManifest, session_state: Any | None, options: ReportExportOptions) -> None:
    add_report_heading(document, "Draft Figures", level=1)
    if not options.include_figures:
        document.add_paragraph("Figures were not included in this draft report per export options.")
        return
    export_ready_items = [item for item in manifest.figure_export_items if item.export_ready]
    if not export_ready_items:
        document.add_paragraph("No export-ready figures are currently available.")
        return
    for index, item in enumerate(export_ready_items, start=1):
        add_report_heading(document, f"Figure {index} - {item.title}", level=2)
        if item.selected_context:
            document.add_paragraph(f"Selected context: {item.selected_context}")
        document.add_paragraph(f"Figure source: {item.source}")
        if item.limitations:
            document.add_paragraph("Limitations: " + "; ".join(item.limitations))
        fig, figure_warnings = build_exportable_figure(item.figure_key, session_state or {}, manifest.figure_context)
        if fig is None:
            document.add_paragraph("[Figure not embedded: source data is unavailable.]")
            for warning in figure_warnings:
                document.add_paragraph(str(warning), style="List Bullet")
            continue
        png_bytes, png_warnings = plotly_figure_to_png_bytes(fig)
        if png_bytes:
            document.add_picture(BytesIO(png_bytes), width=Inches(6.4))
            document.add_paragraph(f"Figure {index} - {item.title}")
        else:
            document.add_paragraph("[Figure not embedded: PNG export unavailable. HTML export remains available from the application.]")
            for warning in [*figure_warnings, *png_warnings]:
                document.add_paragraph(str(warning), style="List Bullet")



def _add_railway_u_girder_sls_report_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add REPORT.RAIL.UGIRDER1 guarded SLS report tables when applicable."""

    package = build_railway_u_girder_sls_report_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Railway U-Girder SLS Engineering Review", level=1)
    document.add_paragraph(
        "This section is generated for the Railway U-Girder staged SLS workflow. "
        "It is an engineering-review report section only and is not a final design certification."
    )
    document.add_paragraph(
        "Decision wording is limited to Preview PASS / REVIEW. Excluded checks include transfer/development length, "
        "anchorage/end-zone bursting, lifting hardware, creep/shrinkage redistribution, ULS coupling, and final code-certified design checks."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")

    table_sequence = [
        ("Closeout Status", package.closeout_status),
        ("Scope and Guardrails", package.scope),
        ("Geometry Summary", package.geometry_summary),
        ("Material and Stage Settings", package.material_stage_settings),
        ("Stage Quantities", package.stage_quantities),
        ("Prestress / Debonding Summary", package.prestress_debonding_summary),
        ("Staged SLS Governing Stress Rows", package.stage_governing),
        ("Staged SLS Limit Governing Rows", package.limit_governing),
        ("Final Service Governing Rows", package.final_service_governing),
        ("SLS Decision Summary", package.decision_summary),
        ("Service Multi-Fiber Summary", package.service_multifiber_summary),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)



def _add_generic_precast_lifting_report_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add GIRDER.LIFT.REPORT1 generic precast lifting report tables."""

    package = build_generic_precast_lifting_report_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Generic Precast Lifting Stage Stress Check", level=1)
    document.add_paragraph(
        "This section is generated by GIRDER.LIFT.REPORT1 for non-Railway precast girder members. "
        "It reports the existing engineering-review lifting-stage preview and does not add a new lifting solver."
    )
    document.add_paragraph(
        "Load basis is limited to individual precast unit self-weight multiplied by lifting impact factor. "
        "Wet slab/topping, barrier, wearing surface, SDL/LL, and building service loads are excluded. "
        "Lifting insert/local hardware, rigging, anchorage, transfer/development length, and final code certification remain outside this report section."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")

    table_sequence = [
        ("Scope and Guardrails", package.scope),
        ("Lifting Settings", package.settings),
        ("Load Basis", package.load_basis),
        ("Station Stress Rows", package.station_stress_rows),
        ("Governing Rows", package.governing_rows),
        ("Closeout Guard", package.closeout_guard),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)

def _add_column_pier_vt_report_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add stored Column/Pier Shear + Torsion report-preview tables when available."""

    package = build_column_pier_vt_report_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Column/Pier Shear + Torsion Strength Gate", level=1)
    document.add_paragraph(
        "This section mirrors the stored Analysis > ULS Strength > Shear + Torsion result. "
        "It reports the strength gate, controlling cause, compact results, audit detail, and scope guard without rerunning solvers during report export."
    )
    document.add_paragraph(
        "Seismic confinement/detailing is a separate review item and is intentionally reported as warning / amber semantics, not as a V+T strength failure. "
        "Prestressed/general-procedure V+T, hoop anchorage/hooks, lap splices, shop-drawing detailing, and final code certification remain outside this report-preview gate."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")
    table_sequence = [
        ("Column/Pier V+T Summary", package.summary),
        ("Column/Pier V+T Compact Results", package.results),
        ("Column/Pier V+T Audit Details", package.audit),
        ("Column/Pier V+T Scope Guard", package.scope_guard),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)


def _add_railway_u_girder_uls_framework_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add ULS.RAIL.UGIRDER1 guarded ULS framework tables when applicable."""

    package = build_railway_u_girder_uls_framework_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Railway U-Girder ULS Strength Check Framework", level=1)
    document.add_paragraph(
        "This section is generated by ULS.RAIL.UGIRDER1. It provides ULS demand traceability, "
        "code-basis guardrails, flexure evidence, PSC shear route evidence, torsion / V+T guard evidence, prestress transfer/development evidence, anchorage/end-zone evidence, and a check-readiness matrix. It is not final code-certified design "
        "and is not an engineer certification."
    )
    document.add_paragraph(
        "Allowed wording is limited to framework-ready / engineering-review / guarded preview. "
        "Do not use code-certified pass wording or final-design pass wording from this milestone."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")

    table_sequence = [
        ("ULS Closeout Boundary", package.closeout_boundary),
        ("ULS Code Basis", package.code_basis),
        ("ULS Demand Summary", package.demand_summary),
        ("ULS Flexure Calculation Evidence", package.flexure_evidence),
        ("ULS PSC Shear Route Evidence", package.shear_evidence),
        ("ULS Torsion / V+T Guard Evidence", package.torsion_vt_guard),
        ("Prestress Transfer / Development Evidence", package.prestress_development_evidence),
        ("Anchorage / End-Zone Evidence", package.anchorage_end_zone_evidence),
        ("ULS Check Matrix", package.check_matrix),
        ("ULS Future Checks", package.future_checks),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)


def _add_railway_u_girder_release_closeout_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add RELEASE.RAIL.UGIRDER1 final engineering-review closeout tables."""

    package = build_railway_u_girder_release_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Railway U-Girder Release Closeout", level=1)
    document.add_paragraph(
        "This section is generated by RELEASE.RAIL.UGIRDER1. It freezes the current Railway U-Girder "
        "SLS + guarded ULS evidence package as an engineering-review release baseline. It adds no new UI, "
        "performs no solver recalculation, and is not final code-certified design."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")
    table_sequence = [
        ("Release Manifest", package.release_manifest),
        ("Release Readiness", package.release_readiness),
        ("Final Claim Guard", package.final_claim_guard),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)


def _add_railway_u_girder_final_design_check_section(
    document: DocumentObject,
    session_state: Any | None,
    options: ReportExportOptions,
) -> None:
    """Add FINAL.RAIL.UGIRDER1 final design-check evidence tables."""

    package = build_railway_u_girder_final_design_check_package(session_state or {})
    if not package.available:
        return
    add_report_heading(document, "Railway U-Girder Final Design-Check Evidence", level=1)
    document.add_paragraph(
        "This section is generated by FINAL.RAIL.UGIRDER1. It consolidates staged SLS evidence, "
        "guarded ULS flexure/shear/torsion-V+T evidence, prestress transfer/development evidence, "
        "anchorage/end-zone evidence, release traceability, and QA guardrails into a final software "
        "design-check evidence package. It is not legal engineer certification and must not be described "
        "as Final Code-Certified Design Complete without Engineer-of-Record approval."
    )
    document.add_paragraph(
        "Allowed wording: Final Design-Check Evidence Package - Complete. Blocked wording: certified approval PASS wording "
        "or software-only final certified design approval."
    )
    for warning in package.warnings:
        document.add_paragraph(str(warning), style="List Bullet")
    table_sequence = [
        ("Final Design-Check Manifest", package.final_manifest),
        ("Final Prerequisite Matrix", package.prerequisite_matrix),
        ("Final Certification Boundary", package.certification_boundary),
        ("Final Handoff", package.final_handoff),
    ]
    for title, dataframe in table_sequence:
        dataframe_to_word_table(document, dataframe, title=title, max_rows=options.max_table_rows)

def _add_generation_notes(document: DocumentObject, manifest: ReportManifest) -> None:
    add_report_heading(document, "Report Generation Notes", level=1)
    rows = [
        ("Generated status", manifest.generated_status),
        ("Data source", "Stored session results and report manifest"),
        ("Solver recalculation", "No solver recalculation during report export"),
        ("Word report status", "Draft engineering report"),
        ("PDF export", "Future work"),
        ("Calculation changes", "Report generation does not change PMM/SLS calculations"),
    ]
    add_key_value_table(document, rows)


def build_draft_word_report(
    manifest: ReportManifest,
    session_state: Any | None = None,
    options: ReportExportOptions | None = None,
) -> bytes:
    options = options or ReportExportOptions()
    document = Document()
    apply_document_defaults(document)
    apply_heading_styles(document)
    add_report_footer_note(document)
    _add_cover(document, manifest)
    _add_executive_summary(document, manifest)
    _add_analysis_scope(document, manifest)
    _add_railway_u_girder_sls_report_section(document, session_state, options)
    _add_generic_precast_lifting_report_section(document, session_state, options)
    _add_column_pier_vt_report_section(document, session_state, options)
    _add_railway_u_girder_uls_framework_section(document, session_state, options)
    _add_railway_u_girder_release_closeout_section(document, session_state, options)
    _add_railway_u_girder_final_design_check_section(document, session_state, options)

    add_report_heading(document, "Result Traceability Snapshot", level=1)
    dataframe_to_word_table(document, result_traceability_snapshot_to_dataframe(manifest.traceability_snapshot), max_rows=options.max_table_rows)

    add_report_heading(document, "Report Readiness", level=1)
    dataframe_to_word_table(document, report_readiness_to_dataframe(manifest.readiness_summary), max_rows=options.max_table_rows)

    add_warning_box_or_section(document, manifest.engineering_warnings, "Engineering Warnings")
    add_limitations_section(document, manifest.engineering_limitations)
    add_unit_conventions_section(document, manifest.unit_conventions)

    add_report_heading(document, "Available Report Tables", level=1)
    dataframe_to_word_table(document, report_tables_to_dataframe(manifest.tables), max_rows=options.max_table_rows)

    add_report_heading(document, "Available Report Figures", level=1)
    figure_summary = report_figure_export_items_to_dataframe(manifest.figure_export_items)[
        ["Figure Key", "Title", "Available", "Export Ready", "Selected Context", "Warning"]
    ]
    dataframe_to_word_table(document, figure_summary, max_rows=options.max_table_rows)

    _add_figures(document, manifest, session_state, options)
    _add_generation_notes(document, manifest)

    if options.include_appendices:
        add_report_heading(document, "Appendices", level=1)
        dataframe_to_word_table(document, report_sections_to_dataframe(manifest.sections), title="Report Section Plan", max_rows=options.max_table_rows)
        if options.include_full_registries:
            dataframe_to_word_table(document, report_tables_to_dataframe(manifest.tables), title="Full Report Table Registry", max_rows=options.max_table_rows)
            dataframe_to_word_table(document, report_figure_export_items_to_dataframe(manifest.figure_export_items), title="Full Report Figure Registry", max_rows=options.max_table_rows)
        if options.include_full_terminology:
            add_terminology_section(document, manifest.terminology)
        dataframe_to_word_table(document, report_manifest_to_summary_dataframe(manifest), title="Raw Manifest Summary", max_rows=options.max_table_rows)

    document.add_paragraph("PDF export remains future work.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def save_draft_word_report(
    manifest: ReportManifest,
    output_path: str,
    session_state: Any | None = None,
    options: ReportExportOptions | None = None,
) -> None:
    Path(output_path).write_bytes(build_draft_word_report(manifest, session_state=session_state, options=options))
