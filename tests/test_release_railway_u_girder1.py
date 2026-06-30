from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.reporting import (
    RAILWAY_UGIRDER_RELEASE_MILESTONE,
    RAILWAY_UGIRDER_RELEASE_STATUS,
    build_draft_word_report,
    build_railway_u_girder_release_package,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_prestress_development1_railway_u_girder import _state_with_debonded_row


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_release_rail_ugirder1_builds_closeout_manifest_without_certification_claim() -> None:
    state = _state_with_debonded_row()
    package = build_railway_u_girder_release_package(state)

    assert package.available is True
    assert package.status == RAILWAY_UGIRDER_RELEASE_STATUS
    assert package.release_manifest.empty is False
    assert package.release_readiness.empty is False
    assert package.final_claim_guard.empty is False
    assert RAILWAY_UGIRDER_RELEASE_MILESTONE in set(package.release_manifest["Status"])
    assert "NO NEW UI" in set(package.release_manifest["Status"])
    assert "NO SOLVER CHANGE" in set(package.release_manifest["Status"])
    assert "NOT CERTIFIED" in set(package.release_manifest["Status"])
    assert "Final design certification" in set(package.release_readiness["Area"])
    assert package.release_readiness.set_index("Area").loc["Final design certification", "Release status"] == "NOT CERTIFIED"
    assert "BLOCKED" in set(package.final_claim_guard["Guard status"])
    assert "final code-certified design" not in package.status.casefold()


def test_release_rail_ugirder1_report_registry_and_word_export_include_release_closeout() -> None:
    state = _state_with_debonded_row()
    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}

    assert by_key["railway_u_girder_release_manifest"].available is True
    assert by_key["railway_u_girder_release_readiness"].available is True
    assert by_key["railway_u_girder_final_claim_guard"].available is True
    assert "not final code-certified design" in by_key["railway_u_girder_release_manifest"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Railway U-Girder Release Closeout" in text
    assert "Release Manifest" in text
    assert "Release Readiness" in text
    assert "Final Claim Guard" in text
    assert RAILWAY_UGIRDER_RELEASE_STATUS in text
    assert "NO NEW UI" in text
    assert "NO SOLVER CHANGE" in text
    assert "not final code-certified design" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_release_rail_ugirder1_docs_lock_no_ui_no_solver_closeout_boundary() -> None:
    release_source = Path("concrete_pmm_pro/reporting/railway_u_girder_release.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    ui_source = Path("concrete_pmm_pro/ui/analysis_page.py").read_text(encoding="utf-8")
    doc = Path("docs/design/release_rail_ugirder1.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "RELEASE.RAIL.UGIRDER1" in release_source
    assert "RAILWAY_UGIRDER_RELEASE_WARNING" in release_source
    assert "railway_u_girder_release_manifest" in report_tables_source
    assert "Railway U-Girder Release Closeout" in word_source
    assert "RELEASE.RAIL.UGIRDER1" not in ui_source
    assert "not final code-certified design" in doc
    assert "No SLS solver equations" in doc
    assert "does not add new UI" in doc
    assert "### RELEASE.RAIL.UGIRDER1" in readme
