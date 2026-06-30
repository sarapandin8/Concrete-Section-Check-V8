from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document

from concrete_pmm_pro.core.analysis import AnalysisModeSettings
from concrete_pmm_pro.geometry.generators import rectangle
from concrete_pmm_pro.reporting.generic_precast_lifting_report import (
    GENERIC_PRECAST_LIFTING_TABLE_KEYS,
    build_generic_precast_lifting_report_package,
    generic_precast_lifting_report_tables_to_dataframe,
    is_generic_precast_lifting_report_context,
)
from concrete_pmm_pro.reporting.report_manifest import build_report_manifest
from concrete_pmm_pro.reporting.report_tables import collect_available_report_tables
from concrete_pmm_pro.reporting.word_export import ReportExportOptions, build_draft_word_report

REPO_ROOT = Path(__file__).resolve().parents[1]
WORD_EXPORT_SOURCE = (REPO_ROOT / "concrete_pmm_pro" / "reporting" / "word_export.py").read_text(encoding="utf-8")
REPORT_TABLES_SOURCE = (REPO_ROOT / "concrete_pmm_pro" / "reporting" / "report_tables.py").read_text(encoding="utf-8")
REPORT_SECTIONS_SOURCE = (REPO_ROOT / "concrete_pmm_pro" / "reporting" / "report_sections.py").read_text(encoding="utf-8")


def _generic_precast_session() -> dict[str, object]:
    geometry = rectangle(900.0, 1500.0).model_copy(
        update={
            "name": "Precast I-Girder QA",
            "metadata": {"preset_key": "parametric_i_girder", "display_name": "Precast I-Girder"},
        }
    )
    strand_table = pd.DataFrame(
        [
            {
                "Active": True,
                "Group ID": "B1",
                "No. Strands": 8,
                "Area/Strand_mm2": 98.7,
                "y_mm_from_bottom": 125.0,
                "Left debond m": 1.0,
                "Right debond m": 1.0,
                "Pe_transfer/strand_kN": 135.0,
                "Pe_construction/strand_kN": 130.0,
                "Pe_eff_final/strand_kN": 115.0,
            }
        ]
    )
    return {
        "analysis_mode_settings": AnalysisModeSettings(member_type="beam_girder"),
        "section_geometry": geometry,
        "section_preset_key": "parametric_i_girder",
        "section_preset_name": "Precast I-Girder",
        "section_parameters": {},
        "girder_strand_layout_table": strand_table,
        "beam_girder_system_settings": {
            "span_length_m": 24.0,
            "girder_spacing_m": 2.5,
            "number_of_girders": 6,
            "concrete_unit_weight_kN_m3": 24.0,
            "lifting_point_ratio": 0.20,
            "lifting_impact_factor": 1.15,
        },
        "beam_girder_sls_auto_load_settings": {
            "include_construction_wet_topping": True,
            "include_service_barrier_sidewalk": True,
            "include_service_wearing_surface": True,
            "include_service_other_sdl": True,
            "other_sdl_area_load_kN_m2": 10.0,
        },
    }


def test_girder_lift_report1_context_detects_generic_and_excludes_railway() -> None:
    session = _generic_precast_session()
    assert is_generic_precast_lifting_report_context(session)

    rail = dict(session)
    rail["section_preset_key"] = "railway_u_girder"
    rail["section_preset_name"] = "Railway U-Girder"
    rail["section_geometry"] = session["section_geometry"].model_copy(
        update={"metadata": {"preset_key": "railway_u_girder", "display_name": "Railway U-Girder"}}
    )
    assert not is_generic_precast_lifting_report_context(rail)


def test_girder_lift_report1_package_generates_settings_load_stress_and_governing_tables() -> None:
    package = build_generic_precast_lifting_report_package(_generic_precast_session())

    assert package.available
    assert set(package.tables()) == set(GENERIC_PRECAST_LIFTING_TABLE_KEYS)
    assert "individual precast unit" in package.scope.to_string().casefold()
    component_rows = package.load_basis.loc[package.load_basis["Component"] != "Excluded by guardrail", "Component"].astype(str).tolist()
    assert all("Wet slab" not in item for item in component_rows)
    assert any("SDL" in str(item) for item in package.load_basis["w_kN/m per girder"].tolist())  # only in the excluded-by-guardrail row
    assert "Precast unit self-weight × lifting IF" in package.load_basis.to_string()
    assert {"Station x (m)", "Auto Mx (kN-m)", "Auto Vy (kN)", "Pe transfer (kN)", "Top total stress (MPa)", "Bottom total stress (MPa)"}.issubset(package.station_stress_rows.columns)
    assert not package.governing_rows.empty
    assert set(package.governing_rows["Demand"]) == {"Governing compression", "Governing tension"}
    assert (package.station_stress_rows["Pe transfer (kN)"] > 0.0).any()


def test_girder_lift_report1_manifest_table_registry_and_section_plan_include_generic_lifting() -> None:
    session = _generic_precast_session()
    tables = collect_available_report_tables(session)
    table_keys = {item.table_key for item in tables if item.available}
    for key in GENERIC_PRECAST_LIFTING_TABLE_KEYS:
        assert key in table_keys

    manifest = build_report_manifest(session)
    assert any(section.section_id == "generic_precast_lifting_stage_report" for section in manifest.sections)
    registry = generic_precast_lifting_report_tables_to_dataframe(build_generic_precast_lifting_report_package(session))
    assert len(registry) == len(GENERIC_PRECAST_LIFTING_TABLE_KEYS)
    assert registry["Available"].all()


def test_girder_lift_report1_word_export_contains_generic_lifting_section() -> None:
    session = _generic_precast_session()
    manifest = build_report_manifest(session)
    report_bytes = build_draft_word_report(
        manifest,
        session,
        options=ReportExportOptions(include_figures=False, include_appendices=False, max_table_rows=12),
    )
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(report_bytes)).paragraphs)

    assert "Generic Precast Lifting Stage Stress Check" in text
    assert "individual precast unit self-weight" in text
    assert "Lifting insert/local hardware" in text


def test_girder_lift_report1_source_integration_guards() -> None:
    assert "_add_generic_precast_lifting_report_section" in WORD_EXPORT_SOURCE
    assert "build_generic_precast_lifting_report_package" in WORD_EXPORT_SOURCE
    assert "generic_precast_lifting_station_stress_rows" in REPORT_TABLES_SOURCE
    assert "generic_precast_lifting_stage_report" in REPORT_SECTIONS_SOURCE

ANALYSIS_SOURCE = (REPO_ROOT / "concrete_pmm_pro" / "ui" / "analysis_page.py").read_text(encoding="utf-8")


def test_girder_lift_report1_analysis_report_qa_panel_is_wired() -> None:
    assert "_render_generic_precast_lifting_report_preview_panel" in ANALYSIS_SOURCE
    assert "Download Generic Precast Lifting Report Table Registry CSV" in ANALYSIS_SOURCE
    assert "generic_precast_lifting_report_package_available" in ANALYSIS_SOURCE
