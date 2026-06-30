from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.analysis.railway_u_girder_uls import (
    RAILWAY_UGIRDER_ULS_TORSION_VT_GUARD_STATUS,
    build_railway_u_girder_uls_framework_package,
    railway_u_girder_uls_torsion_vt_guard_dataframe,
)
from concrete_pmm_pro.core.models import Rebar
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_uls_railway_u_girder3 import _state_with_shear_reinforcement


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def _state_with_torsion_rebar():
    state = _state_with_shear_reinforcement()
    state["section_has_ordinary_rebar"] = True
    state["beam_uls_loads_table"][0]["Tu"] = 20000.0
    # Provide ordinary longitudinal bars as the Al source of truth. Coordinates
    # are deliberately perimeter-like but the torsion guard only counts active
    # ordinary rebar area; drawing-level distribution remains engineer review.
    bars: list[Rebar] = []
    for i in range(42):
        x = -2600.0 + (i % 14) * 400.0
        y = -725.0 if i < 14 else (725.0 if i < 28 else 0.0)
        bars.append(Rebar(x_mm=x, y_mm=y, diameter_mm=32.0, material_name="SD50", label=f"Al{i+1}"))
    state["rebars"] = bars
    return state


def test_uls_rail_ugirder4_builds_guarded_torsion_vt_evidence() -> None:
    state = _state_with_torsion_rebar()
    guard = railway_u_girder_uls_torsion_vt_guard_dataframe(state)

    assert not guard.empty
    assert "ULS torsion / V+T guard" in set(guard["Check"])
    assert RAILWAY_UGIRDER_ULS_TORSION_VT_GUARD_STATUS in set(guard["Evidence status"])
    support = guard.set_index("Case").loc["Strength I - support"]
    assert support["Demand Tu (kN-m)"] == 20000.0
    assert support["Demand Vuy (kN)"] == 980.0
    assert support["φTcr (kN-m)"] > 0.0
    assert support["φ"] == 0.90
    assert support["Ao (mm2)"] > 0.0
    assert support["At/s provided (mm2/m)"] > 0.0
    assert support["Al provided (mm2)"] > 0.0
    assert support["Shear D/C"] >= 0.0
    assert support["V+T interaction index"] >= 0.0
    assert "AASHTO LRFD-compatible" in str(support["Code basis"])
    assert "linear V+T review index" in str(support["Method"])
    assert "No code-certified" in str(support["Blocked final claim"])
    assert "not final code-certified" in str(support["Notes"])


def test_uls_rail_ugirder4_missing_ordinary_rebar_does_not_create_false_vt_pass() -> None:
    state = _state_with_shear_reinforcement()
    state["section_has_ordinary_rebar"] = True
    state["beam_uls_loads_table"][0]["Tu"] = 20000.0
    state["rebars"] = []
    guard = railway_u_girder_uls_torsion_vt_guard_dataframe(state)

    assert not guard.empty
    support = guard.set_index("Case").loc["Strength I - support"]
    # If torsion is below threshold this row may be threshold-only; otherwise it
    # must not report engineering-review PASS without Al.
    if support["Threshold status"] != "BELOW THRESHOLD":
        assert support["Status"] in {"LAYOUT REQUIRED", "Engineering Review FAIL", "REVIEW"}
        assert support["Longitudinal Al status"] in {"LAYOUT REQUIRED", "FAIL", "NOT CHECKED"}
    assert "Code-Certified PASS" not in " ".join(guard.astype(str).stack().tolist())


def test_uls_rail_ugirder4_package_report_registry_and_word_include_torsion_vt_guard_without_certified_claim() -> None:
    state = _state_with_torsion_rebar()
    package = build_railway_u_girder_uls_framework_package(state)

    assert "railway_u_girder_uls_torsion_vt_guard" in package.tables()
    assert not package.torsion_vt_guard.empty
    assert any("ULS.RAIL.UGIRDER4" in warning for warning in package.warnings)

    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}
    assert by_key["railway_u_girder_uls_torsion_vt_guard"].available is True
    assert "not final code-certified" in by_key["railway_u_girder_uls_torsion_vt_guard"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "ULS Torsion / V+T Guard Evidence" in text
    assert RAILWAY_UGIRDER_ULS_TORSION_VT_GUARD_STATUS in text
    assert "linear V+T review index" in text
    assert "not final code-certified design" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_uls_rail_ugirder4_source_markers_and_docs_lock_torsion_vt_boundary() -> None:
    module_source = Path("concrete_pmm_pro/analysis/railway_u_girder_uls.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    doc = Path("docs/design/uls_rail_ugirder4.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "ULS.RAIL.UGIRDER4" in module_source
    assert "RAILWAY_UGIRDER_ULS_TORSION_VT_GUARD_WARNING" in module_source
    assert "railway_u_girder_uls_torsion_vt_guard" in report_tables_source
    assert "ULS Torsion / V+T Guard Evidence" in word_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
    assert "### ULS.RAIL.UGIRDER4" in readme
