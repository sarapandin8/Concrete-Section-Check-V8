from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.reporting import (
    RAILWAY_UGIRDER_CLOSEOUT_STATUS,
    build_draft_word_report,
    build_railway_u_girder_sls_report_package,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_report_railway_u_girder1 import _session_state


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_closeout_rail_ugirder1_package_contains_explicit_closeout_table() -> None:
    package = build_railway_u_girder_sls_report_package(_session_state())

    assert package.available is True
    assert package.closeout_status.empty is False
    assert package.closeout_status.iloc[0]["Status"] == RAILWAY_UGIRDER_CLOSEOUT_STATUS
    assert "not final code-certified" in " ".join(package.closeout_status["Evidence / Boundary"].astype(str)).casefold()
    assert "NOT CERTIFIED" in set(package.closeout_status["Status"])
    assert "READY FOR ENGINEERING REVIEW" in set(package.closeout_status["Status"])
    assert "railway_u_girder_closeout_status" in package.tables()


def test_closeout_rail_ugirder1_report_registry_and_word_export_include_closeout_status() -> None:
    state = _session_state()
    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}

    assert by_key["railway_u_girder_closeout_status"].available is True
    assert by_key["railway_u_girder_closeout_status"].title == "Railway U-Girder Closeout Status"

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Railway U-Girder SLS Engineering Review" in text
    assert "Closeout Status" in text
    assert RAILWAY_UGIRDER_CLOSEOUT_STATUS in text
    assert "NOT CERTIFIED" in text
    assert "Do not use Final Design PASS" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_closeout_rail_ugirder1_docs_and_readme_lock_closeout_boundary() -> None:
    readme = Path("README.md").read_text(encoding="utf-8")
    doc = Path("docs/design/closeout_rail_ugirder1.md").read_text(encoding="utf-8")
    module_source = Path("concrete_pmm_pro/reporting/railway_u_girder_report.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")

    assert readme.startswith("### CLOSEOUT.RAIL.UGIRDER1")
    assert "does **not** mean final code-certified design" in doc
    assert "No solver equations" in doc
    assert "railway_u_girder_closeout_status_dataframe" in module_source
    assert "Closeout Status" in word_source
    assert "Final code-certified design complete" in doc
