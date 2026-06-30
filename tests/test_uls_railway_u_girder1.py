from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.analysis.railway_u_girder_uls import (
    RAILWAY_UGIRDER_ULS_FRAMEWORK_STATUS,
    active_railway_u_girder_uls_demand_dataframe,
    build_railway_u_girder_uls_framework_package,
    is_railway_u_girder_uls_context,
)
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_report_railway_u_girder1 import _session_state


def _uls_session_state():
    state = dict(_session_state())
    state["beam_uls_loads_table"] = [
        {"Active": True, "Station x (m)": 0.0, "Case Name": "Strength I - support", "Mux": 0.0, "Vuy": 980.0, "Tu": 45.0, "Muy": 0.0, "Vux": 0.0, "Nu": 0.0, "Note": "support resultant"},
        {"Active": True, "Station x (m)": 5.0, "Case Name": "Strength I - midspan", "Mux": 2150.0, "Vuy": 0.0, "Tu": 30.0, "Muy": 0.0, "Vux": 0.0, "Nu": 0.0, "Note": "midspan flexure"},
        {"Active": False, "Station x (m)": 10.0, "Case Name": "inactive", "Mux": 9999.0, "Vuy": 9999.0, "Tu": 9999.0, "Muy": 0.0, "Vux": 0.0, "Nu": 0.0, "Note": "inactive"},
    ]
    return state


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_uls_rail_ugirder1_context_and_active_uls_demands_are_routed_from_loads() -> None:
    state = _uls_session_state()

    assert is_railway_u_girder_uls_context(state) is True
    active = active_railway_u_girder_uls_demand_dataframe(state)

    assert len(active) == 2
    assert set(active["Case Name"]) == {"Strength I - support", "Strength I - midspan"}
    assert active["Mux"].max() == 2150.0
    assert active["Vuy"].max() == 980.0
    assert active["Tu"].max() == 45.0


def test_uls_rail_ugirder1_package_builds_guarded_framework_tables_without_certified_claim() -> None:
    package = build_railway_u_girder_uls_framework_package(_uls_session_state())

    assert package.available is True
    assert package.status == RAILWAY_UGIRDER_ULS_FRAMEWORK_STATUS
    assert "not final code-certified" in " ".join(package.warnings).casefold()
    assert "railway_u_girder_uls_check_matrix" in package.tables()
    assert "NOT CERTIFIED" in set(package.closeout_boundary["Status"])
    assert "FRAMEWORK READY" in set(package.check_matrix["Current Framework Status"])
    assert "GUARDED PREVIEW" in set(package.check_matrix["Current Framework Status"])
    assert "FUTURE WORK" in set(package.check_matrix["Current Framework Status"])
    assert not package.demand_summary.empty
    assert package.demand_summary.set_index("Demand Item").loc["Active ULS station rows", "Value"] == 2.0
    assert package.demand_summary.set_index("Demand Item").loc["Peak |Mux|", "Value"] == 2150.0


def test_uls_rail_ugirder1_report_registry_and_word_export_include_guarded_uls_section() -> None:
    state = _uls_session_state()
    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}

    assert by_key["railway_u_girder_uls_check_matrix"].available is True
    assert by_key["railway_u_girder_uls_check_matrix"].warning
    assert "not final code-certified" in by_key["railway_u_girder_uls_check_matrix"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Railway U-Girder ULS Strength Check Framework" in text
    assert "ULS Check Matrix" in text
    assert "ULS Demand Summary" in text
    assert RAILWAY_UGIRDER_ULS_FRAMEWORK_STATUS in text
    assert "not final code-certified design" in text
    assert "Engineer-of-Record" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_uls_rail_ugirder1_source_markers_and_docs_lock_framework_boundary() -> None:
    module_source = Path("concrete_pmm_pro/analysis/railway_u_girder_uls.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    ui_source = Path("concrete_pmm_pro/ui/analysis_page.py").read_text(encoding="utf-8")
    doc = Path("docs/design/uls_rail_ugirder1.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "ULS.RAIL.UGIRDER1" in module_source
    assert "RAILWAY_UGIRDER_ULS_REQUIRED_FUTURE_CHECKS" in module_source
    assert "_add_railway_u_girder_uls_framework_section" in word_source
    assert "_render_railway_u_girder_uls_framework_preview_panel" in ui_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
    assert "### ULS.RAIL.UGIRDER1" in readme
