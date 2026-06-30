from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.analysis.railway_u_girder_uls import (
    RAILWAY_UGIRDER_PRESTRESS_DEVELOPMENT_STATUS,
    build_railway_u_girder_uls_framework_package,
    railway_u_girder_prestress_development_evidence_dataframe,
)
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_uls_railway_u_girder1 import _uls_session_state


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def _state_with_debonded_row():
    state = _uls_session_state()
    table = [dict(row) for row in state["girder_strand_layout_table"].to_dict("records")]
    table[0]["Debonded strand nos"] = "1,2,8,9"
    table[0]["Left debond m"] = 2.0
    table[0]["Right debond m"] = 2.0
    state["girder_strand_layout_table"] = table
    return state


def test_prestress_development1_builds_guarded_transfer_development_evidence() -> None:
    state = _state_with_debonded_row()
    evidence = railway_u_girder_prestress_development_evidence_dataframe(state)

    assert not evidence.empty
    assert RAILWAY_UGIRDER_PRESTRESS_DEVELOPMENT_STATUS in set(evidence["Evidence status"])
    first = evidence.set_index("Group ID").loc["L Row 1"]
    assert first["Debonded strands"] == 4
    assert first["Strand diameter (mm)"] == 12.7
    assert first["Transfer length lt (m)"] > 0.0
    assert first["Development length ld (m)"] >= first["Transfer length lt (m)"]
    assert first["Left full-development station (m)"] > first["Left debond (m)"]
    assert first["Right full-development station (m)"] < 10.0 - first["Right debond (m)"]
    assert first["Left development D/C"] >= 0.0
    assert "AASHTO/ACI-compatible" in str(first["Code basis"])
    assert "No code-certified" in str(first["Blocked final claim"])
    assert "no Pe force ramp" in str(first["Notes"]) or "No Pe force ramp" in str(first["Notes"])


def test_prestress_development1_package_report_registry_and_word_include_development_evidence() -> None:
    state = _state_with_debonded_row()
    package = build_railway_u_girder_uls_framework_package(state)

    assert "railway_u_girder_prestress_development_evidence" in package.tables()
    assert not package.prestress_development_evidence.empty
    assert any("PRESTRESS.DEVELOPMENT1" in warning for warning in package.warnings)

    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}
    assert by_key["railway_u_girder_prestress_development_evidence"].available is True
    assert "not final code-certified" in by_key["railway_u_girder_prestress_development_evidence"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Prestress Transfer / Development Evidence" in text
    assert RAILWAY_UGIRDER_PRESTRESS_DEVELOPMENT_STATUS in text
    assert "AASHTO/ACI-compatible" in text
    assert "not final code-certified design" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_prestress_development1_empty_strand_layout_does_not_create_false_pass() -> None:
    state = _uls_session_state()
    state["girder_strand_layout_table"] = []
    evidence = railway_u_girder_prestress_development_evidence_dataframe(state)
    package = build_railway_u_girder_uls_framework_package(state)

    assert evidence.empty
    assert package.prestress_development_evidence.empty
    assert any("No active strand rows" in warning for warning in package.warnings)


def test_prestress_development1_source_markers_and_docs_lock_boundary() -> None:
    module_source = Path("concrete_pmm_pro/analysis/railway_u_girder_uls.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    doc = Path("docs/design/prestress_development1.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "PRESTRESS.DEVELOPMENT1" in module_source
    assert "RAILWAY_UGIRDER_PRESTRESS_DEVELOPMENT_WARNING" in module_source
    assert "railway_u_girder_prestress_development_evidence" in report_tables_source
    assert "Prestress Transfer / Development Evidence" in word_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
    assert "### PRESTRESS.DEVELOPMENT1" in readme
