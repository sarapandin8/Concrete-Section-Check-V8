from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.reporting import (
    RAILWAY_UGIRDER_FINAL_MILESTONE,
    RAILWAY_UGIRDER_FINAL_STATUS,
    build_draft_word_report,
    build_railway_u_girder_final_design_check_package,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_prestress_development1_railway_u_girder import _state_with_debonded_row


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_final_rail_ugirder1_builds_complete_software_evidence_package_with_eor_boundary() -> None:
    state = _state_with_debonded_row()
    package = build_railway_u_girder_final_design_check_package(state)

    assert package.available is True
    assert package.status == RAILWAY_UGIRDER_FINAL_STATUS
    assert package.final_manifest.empty is False
    assert package.prerequisite_matrix.empty is False
    assert package.certification_boundary.empty is False
    assert package.final_handoff.empty is False
    assert RAILWAY_UGIRDER_FINAL_MILESTONE in set(package.final_manifest["Status"])
    assert "NO NEW UI" in set(package.final_manifest["Status"])
    assert "NO SOLVER CHANGE" in set(package.final_manifest["Status"])
    assert "EOR REQUIRED" in set(package.final_manifest["Status"])
    assert "Engineer-of-Record certification" in set(package.prerequisite_matrix["Prerequisite"])
    assert package.prerequisite_matrix.set_index("Prerequisite").loc["Engineer-of-Record certification", "Evidence status"] == "REQUIRED OUTSIDE SOFTWARE"
    assert "BLOCKED WITHOUT EOR" in set(package.certification_boundary["Allowed status"])
    assert any("FINAL.RAIL.UGIRDER1" in warning for warning in package.warnings)


def test_final_rail_ugirder1_registry_and_word_export_include_final_evidence_without_false_certification() -> None:
    state = _state_with_debonded_row()
    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}

    assert by_key["railway_u_girder_final_design_check_manifest"].available is True
    assert by_key["railway_u_girder_final_prerequisite_matrix"].available is True
    assert by_key["railway_u_girder_final_certification_boundary"].available is True
    assert by_key["railway_u_girder_final_handoff"].available is True
    assert "not legal engineer certification" in by_key["railway_u_girder_final_design_check_manifest"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Railway U-Girder Final Design-Check Evidence" in text
    assert "Final Design-Check Manifest" in text
    assert "Final Prerequisite Matrix" in text
    assert "Final Certification Boundary" in text
    assert "Final Handoff" in text
    assert RAILWAY_UGIRDER_FINAL_STATUS in text
    assert "Engineer-of-Record" in text
    assert "BLOCKED WITHOUT EOR" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_final_rail_ugirder1_source_markers_and_docs_lock_no_ui_no_solver_boundary() -> None:
    final_source = Path("concrete_pmm_pro/reporting/railway_u_girder_final.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    ui_source = Path("concrete_pmm_pro/ui/analysis_page.py").read_text(encoding="utf-8")
    doc = Path("docs/design/final_rail_ugirder1.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "FINAL.RAIL.UGIRDER1" in final_source
    assert "RAILWAY_UGIRDER_FINAL_WARNING" in final_source
    assert "railway_u_girder_final_design_check_manifest" in report_tables_source
    assert "Railway U-Girder Final Design-Check Evidence" in word_source
    assert "FINAL.RAIL.UGIRDER1" not in ui_source
    assert "No SLS solver equations" in doc
    assert "No SLS solver equations" in readme
    assert "### FINAL.RAIL.UGIRDER1" in readme
    assert "Final Code-Certified Design Complete" in doc
    assert "BLOCKED WITHOUT EOR" in final_source
