from __future__ import annotations

import base64
from dataclasses import replace
from io import BytesIO

import pandas as pd
import plotly.graph_objects as go
from docx import Document

from concrete_pmm_pro.core.models import LoadCase
from concrete_pmm_pro.reporting import build_report_manifest
from concrete_pmm_pro.reporting.word_export import ReportExportOptions, build_draft_word_report, dataframe_to_word_table
from concrete_pmm_pro.serviceability import ServiceStressPointResult, ServiceabilitySettings, ServiceabilitySummary


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text)


def _simple_serviceability_summary() -> ServiceabilitySummary:
    return ServiceabilitySummary(
        enabled=True,
        settings=ServiceabilitySettings(enabled=True),
        section_properties=None,
        sls_load_cases=[LoadCase(name="SLS Service 1", load_type="SLS")],
        stress_results=[
            ServiceStressPointResult(
                combo_name="SLS Service 1",
                point_name="Top fiber",
                x_mm=0.0,
                y_mm=300.0,
                stress_MPa=1.2,
                total_stress_MPa=1.2,
                external_stress_MPa=1.2,
                prestress_stress_MPa=0.0,
                stress_type="Tension",
                status="PASS",
                utilization=0.6,
                message="Tension stress within allowable limit.",
            )
        ],
        overall_status="PASS",
        governing_combo="SLS Service 1",
        section_basis_used="gross",
    )


def _one_pixel_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )


def test_build_draft_word_report_returns_bytes() -> None:
    report_bytes = build_draft_word_report(build_report_manifest({}))

    assert isinstance(report_bytes, bytes)
    assert report_bytes


def test_report_export_options_default_creation() -> None:
    options = ReportExportOptions()

    assert options.include_appendices is True
    assert options.include_figures is True
    assert options.max_table_rows == 30


def test_build_draft_word_report_accepts_report_export_options() -> None:
    report_bytes = build_draft_word_report(build_report_manifest({}), options=ReportExportOptions(max_table_rows=10))

    assert report_bytes.startswith(b"PK")


def test_build_draft_word_report_returns_valid_docx_zip() -> None:
    report_bytes = build_draft_word_report(build_report_manifest({}))

    assert report_bytes.startswith(b"PK")
    assert Document(BytesIO(report_bytes)).paragraphs


def test_word_report_includes_report_title() -> None:
    manifest = build_report_manifest({"project_name": "Pier Report"})

    assert "Concrete Section Pro Engineering Report" in _doc_text(build_draft_word_report(manifest))


def test_cover_page_includes_draft_disclaimer() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "It is not a final design certification." in text


def test_word_report_uses_guarded_production_preview_scope_wording() -> None:
    report_bytes = build_draft_word_report(build_report_manifest({}))
    text = _doc_text(report_bytes)
    document = Document(BytesIO(report_bytes))
    footer_text = "\n".join(paragraph.text for section in document.sections for paragraph in section.footer.paragraphs)

    assert "ACI RC Flexural PMM may be treated as finalized production-preview only within the validated RC scope" in text
    assert "unsupported routes remain engineering-review items" in text
    assert "finalized production-preview only where explicitly validated" in footer_text
    assert "prototype engineering review only" not in footer_text


def test_word_report_includes_readiness_status() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "Readiness status" in text
    assert "NOT_READY" in text


def test_limitations_section_includes_critical_high_limitations() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "Convex hull fallback for PMM slice envelope" in text
    assert "Directional D/C method" in text


def test_limitations_are_ordered_with_critical_high_first() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert text.find("Convex hull fallback for PMM slice envelope") < text.find("Neutral-axis sweep resolution")


def test_word_report_includes_engineering_limitations_heading() -> None:
    assert "Engineering Limitations" in _doc_text(build_draft_word_report(build_report_manifest({})))


def test_word_report_includes_unit_conventions_heading() -> None:
    assert "Unit Conventions" in _doc_text(build_draft_word_report(build_report_manifest({})))


def test_word_report_includes_terminology_heading() -> None:
    assert "Standard Terminology" in _doc_text(build_draft_word_report(build_report_manifest({})))


def test_word_report_works_with_empty_minimal_manifest() -> None:
    Document(BytesIO(build_draft_word_report(build_report_manifest({}))))


def test_warning_section_handles_empty_warnings() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "No engineering warnings recorded." in text


def test_word_report_works_with_manifest_containing_warnings() -> None:
    manifest = replace(build_report_manifest({}), engineering_warnings=["Review prototype limitations."])
    text = _doc_text(build_draft_word_report(manifest))

    assert "Review prototype limitations." in text


def test_word_report_works_when_figures_unavailable() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "No export-ready figures are currently available." in text


def test_word_report_does_not_fail_when_png_export_unavailable(monkeypatch) -> None:
    def _no_png(_fig):
        return None, ["PNG export requires kaleido. HTML export remains available."]

    monkeypatch.setattr("concrete_pmm_pro.reporting.word_export.plotly_figure_to_png_bytes", _no_png)
    manifest = build_report_manifest({"serviceability_summary": _simple_serviceability_summary()})
    text = _doc_text(build_draft_word_report(manifest, {"serviceability_summary": _simple_serviceability_summary()}))

    assert "Figure not embedded" in text
    assert "PNG export requires kaleido" in text


def test_word_report_embeds_figure_when_png_export_succeeds(monkeypatch) -> None:
    def _png(_fig):
        return _one_pixel_png(), []

    monkeypatch.setattr("concrete_pmm_pro.reporting.word_export.plotly_figure_to_png_bytes", _png)
    session_state = {
        "rc_pmm_result": object(),
        "serviceability_summary": _simple_serviceability_summary(),
        "section_geometry": None,
        "pmm_interaction_surface_figure": go.Figure(data=[go.Scatter3d(x=[0, 1], y=[0, 1], z=[0, 1])]),
    }
    manifest = build_report_manifest(session_state)
    report_bytes = build_draft_word_report(manifest, session_state)
    document = Document(BytesIO(report_bytes))

    assert len(document.inline_shapes) >= 1
    assert "Figure not embedded: PNG export unavailable" not in _doc_text(report_bytes)


def test_include_figures_false_skips_figure_embedding_section() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({}), options=ReportExportOptions(include_figures=False)))

    assert "Figures were not included in this draft report per export options." in text


def test_include_appendices_false_skips_appendix_tables() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({}), options=ReportExportOptions(include_appendices=False)))

    assert "Appendices" not in text
    assert "Report Section Plan" not in text


def test_report_generation_notes_state_no_solver_recalculation() -> None:
    text = _doc_text(build_draft_word_report(build_report_manifest({})))

    assert "Report Generation Notes" in text
    assert "No solver recalculation during report export" in text


def test_dataframe_to_word_table_handles_empty_dataframe() -> None:
    document = Document()
    dataframe_to_word_table(document, pd.DataFrame(), title="Empty Table")

    assert "No data available." in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_dataframe_to_word_table_truncates_large_dataframe() -> None:
    document = Document()
    dataframe_to_word_table(document, pd.DataFrame({"A": range(10)}), title="Large Table", max_rows=3)

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Table truncated for draft report" in text


def test_build_draft_word_report_does_not_trigger_solver_recalculation(monkeypatch) -> None:
    def _raise_if_called(*_args, **_kwargs):
        raise AssertionError("solver should not be called")

    monkeypatch.setattr("concrete_pmm_pro.analysis.pmm_solver.run_rc_pmm_solver", _raise_if_called)
    report_bytes = build_draft_word_report(build_report_manifest({}))

    assert report_bytes.startswith(b"PK")


def test_analysis_page_imports_without_error() -> None:
    import concrete_pmm_pro.ui.analysis_page as analysis_page

    assert hasattr(analysis_page, "render_analysis_page")
