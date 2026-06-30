from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.core.analysis import AnalysisModeSettings
from concrete_pmm_pro.core.models import LoadCase
from concrete_pmm_pro.geometry.generators import railway_u_girder
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_railway_u_girder_sls_report_package,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from concrete_pmm_pro.ui.prestress_page import _default_girder_strand_layout_table


def _railway_geometry():
    return railway_u_girder(
        width_mm=5500,
        depth_mm=1600,
        top_wall_width_mm=600,
        bottom_side_width_mm=650,
        haunch_x_mm=300,
        haunch_y_mm=300,
        h1_step_height_mm=670,
        h2_bottom_opening_mm=305,
        h3_floor_side_thickness_mm=395,
        h4_floor_center_thickness_mm=450,
    )


def _stage_settings():
    return {
        "web_fc_MPa": 45.0,
        "web_fci_MPa": 36.0,
        "slab_fc_MPa": 35.0,
        "concrete_unit_weight_kN_m3": 24.0,
        "support_condition": "Simply supported",
        "construction_method": "Case B - wet slab carried by precast webs",
        "wet_slab_distribution_each_web": 0.5,
        "formwork_construction_load_kN_m2": 2.5,
        "lifting_point_ratio": 0.20,
        "lifting_impact_factor": 1.10,
    }


def _session_state():
    geom = _railway_geometry()
    table = _default_girder_strand_layout_table(geom)
    return {
        "project_name": "Railway U-Girder Report Test",
        "analysis_mode_settings": AnalysisModeSettings(member_type="beam_girder", analysis_workflow="bridge_beam_girder"),
        "section_preset_key": "railway_u_girder",
        "section_preset_name": "Railway U-Girder",
        "section_geometry": geom,
        "section_parameters": geom.metadata.get("parameters", {}),
        "girder_strand_layout_table": table,
        "railway_u_girder_stage_settings": _stage_settings(),
        "beam_girder_system_settings": {
            "girder_system": "Railway U-Girder",
            "span_length_m": 10.0,
            "concrete_unit_weight_kN_m3": 24.0,
        },
        "load_cases": [LoadCase(name="SLS-final", Pu_N=0.0, Mux_Nmm=700_000_000.0, load_type="SLS", active=True)],
    }


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_report_rail_ugirder1_package_builds_guarded_tables() -> None:
    package = build_railway_u_girder_sls_report_package(_session_state())

    assert package.available is True
    assert "not final code-certified" in package.status
    assert set(package.decision_summary["Check stage"]) == {"Transfer", "Lifting", "Wet slab casting", "Final service"}
    assert "Preview PASS" in set(package.decision_summary["Decision"])
    assert "REVIEW" in set(package.decision_summary["Decision"]) or "Preview PASS" in set(package.decision_summary["Decision"])
    assert package.material_stage_settings.set_index("Setting").loc["Precast web f'ci", "Value"] == 36.0
    assert set(package.service_multifiber_summary["Fiber"]) == {
        "Top web fiber",
        "Bottom web fiber",
        "CIP slab top fiber",
        "CIP slab bottom fiber",
    }
    assert set(package.service_multifiber_summary["Concrete component"]) == {"Web", "CIP slab"}


def test_report_rail_ugirder1_registers_tables_for_report_manifest() -> None:
    tables = collect_available_report_tables(_session_state())
    by_key = {table.table_key: table for table in tables}

    assert by_key["railway_u_girder_sls_decision_summary"].available is True
    assert by_key["railway_u_girder_service_multifiber_summary"].available is True
    assert by_key["railway_u_girder_sls_decision_summary"].warning
    assert "not final code-certified" in by_key["railway_u_girder_sls_decision_summary"].warning


def test_report_rail_ugirder1_word_report_includes_guarded_section_and_passes_language_qa() -> None:
    state = _session_state()
    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "Railway U-Girder SLS Engineering Review" in text
    assert "SLS Decision Summary" in text
    assert "Service Multi-Fiber Summary" in text
    assert "It is an engineering-review report section only" in text
    assert "not a final design certification" in text
    assert "final code-certified design checks" in text
    assert "Precast web f'ci" in text
    assert "36" in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_report_rail_ugirder1_source_markers_and_docs() -> None:
    module_source = Path("concrete_pmm_pro/reporting/railway_u_girder_report.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    ui_source = Path("concrete_pmm_pro/ui/analysis_page.py").read_text(encoding="utf-8")
    doc = Path("docs/design/report_rail_ugirder1.md").read_text(encoding="utf-8")

    assert "REPORT.RAIL.UGIRDER1" in module_source
    assert "RAILWAY_UGIRDER_REPORT_EXCLUSIONS" in module_source
    assert "_add_railway_u_girder_sls_report_section" in word_source
    assert "_render_railway_u_girder_report_preview_panel" in ui_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
