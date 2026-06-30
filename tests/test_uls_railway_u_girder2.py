from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.analysis.railway_u_girder_uls import (
    RAILWAY_UGIRDER_ULS_FLEXURE_EVIDENCE_STATUS,
    build_railway_u_girder_uls_framework_package,
    railway_u_girder_uls_flexure_evidence_dataframe,
)
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_uls_railway_u_girder1 import _uls_session_state


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def test_uls_rail_ugirder2_builds_guarded_flexure_calculation_evidence() -> None:
    state = _uls_session_state()
    evidence = railway_u_girder_uls_flexure_evidence_dataframe(state)

    assert not evidence.empty
    assert set(evidence["Check"]) == {"ULS flexure"}
    assert RAILWAY_UGIRDER_ULS_FLEXURE_EVIDENCE_STATUS in set(evidence["Evidence status"])
    assert "Strength I - midspan" in set(evidence["Case"])
    governing = evidence.set_index("Case").loc["Strength I - midspan"]
    assert governing["Demand Mux (kN-m)"] == 2150.0
    assert governing["φMn (kN-m)"] > 0.0
    assert governing["Nominal Mn (kN-m)"] > 0.0
    assert governing["D/C"] >= 0.0
    assert governing["Effective strands"] > 0
    assert "AASHTO LRFD" in str(governing["Code basis"])
    assert "single-material" in str(governing["Concrete basis"])
    assert "No code-certified" in str(governing["Blocked final claim"])
    assert "Engineering Review" in str(governing["Status"])


def test_uls_rail_ugirder2_package_report_registry_and_word_include_flexure_evidence_without_certified_claim() -> None:
    state = _uls_session_state()
    package = build_railway_u_girder_uls_framework_package(state)

    assert "railway_u_girder_uls_flexure_evidence" in package.tables()
    assert not package.flexure_evidence.empty
    assert any("ULS.RAIL.UGIRDER2" in warning for warning in package.warnings)

    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}
    assert by_key["railway_u_girder_uls_flexure_evidence"].available is True
    assert "not final code-certified" in by_key["railway_u_girder_uls_flexure_evidence"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "ULS Flexure Calculation Evidence" in text
    assert RAILWAY_UGIRDER_ULS_FLEXURE_EVIDENCE_STATUS in text
    assert "Engineering Review PASS" in text or "Engineering Review FAIL" in text
    assert "single-material" in text
    assert "not final code-certified design" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_uls_rail_ugirder2_zero_mux_rows_do_not_create_false_flexure_pass() -> None:
    state = _uls_session_state()
    state["beam_uls_loads_table"] = [
        {"Active": True, "Station x (m)": 0.0, "Case Name": "zero support", "Mux": 0.0, "Vuy": 980.0, "Tu": 45.0, "Muy": 0.0, "Vux": 0.0, "Nu": 0.0, "Note": "zero flexure"},
    ]
    evidence = railway_u_girder_uls_flexure_evidence_dataframe(state)
    package = build_railway_u_girder_uls_framework_package(state)

    assert evidence.empty
    assert package.flexure_evidence.empty
    assert any("No nonzero Mux" in warning for warning in package.warnings)


def test_uls_rail_ugirder2_source_markers_and_docs_lock_evidence_boundary() -> None:
    module_source = Path("concrete_pmm_pro/analysis/railway_u_girder_uls.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    doc = Path("docs/design/uls_rail_ugirder2.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "ULS.RAIL.UGIRDER2" in module_source
    assert "RAILWAY_UGIRDER_ULS_FLEXURE_EVIDENCE_WARNING" in module_source
    assert "railway_u_girder_uls_flexure_evidence" in report_tables_source
    assert "ULS Flexure Calculation Evidence" in word_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
    assert "### ULS.RAIL.UGIRDER2" in readme
