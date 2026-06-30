from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from concrete_pmm_pro.analysis.railway_u_girder_uls import (
    RAILWAY_UGIRDER_ULS_SHEAR_EVIDENCE_STATUS,
    build_railway_u_girder_uls_framework_package,
    railway_u_girder_uls_shear_evidence_dataframe,
)
from concrete_pmm_pro.reporting import (
    build_draft_word_report,
    build_report_manifest,
    collect_available_report_tables,
    run_word_report_qa,
)
from tests.test_uls_railway_u_girder1 import _uls_session_state


def _doc_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def _state_with_shear_reinforcement():
    state = _uls_session_state()
    state["beam_girder_shear_reinforcement_table"] = [
        {
            "Active": True,
            "Zone": "Support shear zone",
            "x_start_m": 0.0,
            "x_end_m": 2.0,
            "Bar Size": "DB16",
            "Diameter_mm": 16.0,
            "Legs": 4,
            "Spacing_mm": 150.0,
            "fy_MPa": 390.0,
            "Note": "provided support stirrups",
        },
        {
            "Active": True,
            "Zone": "Midspan shear zone",
            "x_start_m": 2.0,
            "x_end_m": 10.0,
            "Bar Size": "DB12",
            "Diameter_mm": 12.0,
            "Legs": 4,
            "Spacing_mm": 200.0,
            "fy_MPa": 390.0,
            "Note": "provided web stirrups",
        },
    ]
    return state


def test_uls_rail_ugirder3_builds_guarded_psc_shear_route_evidence() -> None:
    state = _state_with_shear_reinforcement()
    evidence = railway_u_girder_uls_shear_evidence_dataframe(state)

    assert not evidence.empty
    assert set(evidence["Check"]) == {"ULS PSC shear"}
    assert RAILWAY_UGIRDER_ULS_SHEAR_EVIDENCE_STATUS in set(evidence["Evidence status"])
    support = evidence.set_index("Case").loc["Strength I - support"]
    assert support["Demand Vuy (kN)"] == 980.0
    assert support["φVn (kN)"] > 0.0
    assert support["φVc (kN)"] > 0.0
    assert support["φVs (kN)"] > 0.0
    assert support["D/C"] >= 0.0
    assert support["bv (mm)"] > 0.0
    assert support["dv (mm)"] > 0.0
    assert support["β"] == 2.0
    assert support["θ (deg)"] == 45.0
    assert support["φ"] == 0.90
    assert "AASHTO LRFD" in str(support["Code basis"])
    assert "Vp = 0" in str(support["PSC shear basis"])
    assert "No code-certified" in str(support["Blocked final claim"])
    assert "Engineering Review" in str(support["Status"])


def test_uls_rail_ugirder3_missing_stirrup_layout_does_not_create_false_shear_pass() -> None:
    state = _uls_session_state()
    evidence = railway_u_girder_uls_shear_evidence_dataframe(state)

    assert not evidence.empty
    support = evidence.set_index("Case").loc["Strength I - support"]
    assert support["Status"] == "LAYOUT REQUIRED"
    assert support["Zone"] == "-"
    assert str(support["Blocked final claim"]).startswith("No code-certified")
    assert "No active shear reinforcement zone" in str(support["Notes"])


def test_uls_rail_ugirder3_package_report_registry_and_word_include_shear_evidence_without_certified_claim() -> None:
    state = _state_with_shear_reinforcement()
    package = build_railway_u_girder_uls_framework_package(state)

    assert "railway_u_girder_uls_shear_evidence" in package.tables()
    assert not package.shear_evidence.empty
    assert any("ULS.RAIL.UGIRDER3" in warning for warning in package.warnings)

    tables = collect_available_report_tables(state)
    by_key = {table.table_key: table for table in tables}
    assert by_key["railway_u_girder_uls_shear_evidence"].available is True
    assert "not final code-certified" in by_key["railway_u_girder_uls_shear_evidence"].warning

    manifest = build_report_manifest(state)
    docx_bytes = build_draft_word_report(manifest, state)
    text = _doc_text(docx_bytes)

    assert "ULS PSC Shear Route Evidence" in text
    assert RAILWAY_UGIRDER_ULS_SHEAR_EVIDENCE_STATUS in text
    assert "Engineering Review PASS" in text or "Engineering Review FAIL" in text
    assert "Vp = 0" in text
    assert "not final code-certified design" in text
    assert "Code-Certified PASS" not in text

    qa = run_word_report_qa(docx_bytes, manifest)
    assert qa.fail_count == 0


def test_uls_rail_ugirder3_source_markers_and_docs_lock_shear_boundary() -> None:
    module_source = Path("concrete_pmm_pro/analysis/railway_u_girder_uls.py").read_text(encoding="utf-8")
    report_tables_source = Path("concrete_pmm_pro/reporting/report_tables.py").read_text(encoding="utf-8")
    word_source = Path("concrete_pmm_pro/reporting/word_export.py").read_text(encoding="utf-8")
    doc = Path("docs/design/uls_rail_ugirder3.md").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")

    assert "ULS.RAIL.UGIRDER3" in module_source
    assert "RAILWAY_UGIRDER_ULS_SHEAR_EVIDENCE_WARNING" in module_source
    assert "railway_u_girder_uls_shear_evidence" in report_tables_source
    assert "ULS PSC Shear Route Evidence" in word_source
    assert "not final code-certified" in doc
    assert "No SLS solver equations" in doc
    assert "### ULS.RAIL.UGIRDER3" in readme
